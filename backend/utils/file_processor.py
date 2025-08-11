"""
파일 처리 유틸리티 함수들
PDF, DOCX 파일 파싱 및 검증 기능 제공
"""

import io
import os
from PyPDF2 import PdfReader
from docx import Document
from config.settings import ALLOWED_EXTENSIONS, MAX_FILE_SIZE


class FileProcessingError(Exception):
    """파일 처리 관련 예외"""
    pass


def validate_file_type(filename):
    """
    파일 타입 검증
    
    Args:
        filename (str): 검증할 파일명
        
    Raises:
        FileProcessingError: 지원하지 않는 파일 형식인 경우
    """
    if not filename:
        raise FileProcessingError("파일명이 제공되지 않았습니다.")
    
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise FileProcessingError(
            f"지원하지 않는 파일 형식입니다: {ext}. "
            f"지원되는 형식: {', '.join(ALLOWED_EXTENSIONS)}"
        )


def validate_file_size(file_stream):
    """
    파일 크기 검증
    
    Args:
        file_stream: 파일 스트림 객체
        
    Raises:
        FileProcessingError: 파일 크기가 제한을 초과하는 경우
    """
    # 현재 위치 저장
    current_position = file_stream.tell()
    
    # 파일 끝으로 이동하여 크기 확인
    file_stream.seek(0, 2)  # 파일 끝으로 이동
    file_size = file_stream.tell()
    
    # 원래 위치로 복귀
    file_stream.seek(current_position)
    
    if file_size > MAX_FILE_SIZE:
        # 사용자 메시지를 프론트에서 그대로 노출할 수 있도록 명확히 표기
        raise FileProcessingError("첨부파일의 용량이 50mb를 초과했습니다.")


def parse_pdf(file_stream):
    """
    PDF 파일을 파싱하여 텍스트 추출
    
    Args:
        file_stream: PDF 파일 스트림
        
    Returns:
        str: 추출된 텍스트
        
    Raises:
        FileProcessingError: PDF 파싱 실패 시
    """
    try:
        # 파일 크기 검증
        validate_file_size(file_stream)
        
        # PDF 파싱
        stream = io.BytesIO(file_stream.read())
        reader = PdfReader(stream)
        
        if len(reader.pages) == 0:
            raise FileProcessingError("PDF 파일에 페이지가 없습니다.")
        
        text = ""
        for page_num, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
            except Exception as e:
                # 개별 페이지 오류는 로깅만 하고 계속 진행
                print(f"페이지 {page_num + 1} 처리 중 오류: {str(e)}")
        
        if not text.strip():
            raise FileProcessingError("PDF에서 텍스트를 추출할 수 없습니다.")
        
        return text.strip()
        
    except FileProcessingError:
        raise
    except Exception as e:
        raise FileProcessingError(f"PDF 파싱 중 오류 발생: {str(e)}")


def parse_docx(file_stream):
    """
    DOCX 파일을 파싱하여 텍스트 추출
    
    Args:
        file_stream: DOCX 파일 스트림
        
    Returns:
        str: 추출된 텍스트
        
    Raises:
        FileProcessingError: DOCX 파싱 실패 시
    """
    try:
        # 파일 크기 검증
        validate_file_size(file_stream)
        
        # DOCX 파싱
        stream = io.BytesIO(file_stream.read())
        doc = Document(stream)
        
        text = ""
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text += paragraph.text + "\n"
        
        # 표(table) 내용도 추출
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text += cell.text + "\n"
        
        if not text.strip():
            raise FileProcessingError("DOCX에서 텍스트를 추출할 수 없습니다.")
        
        return text.strip()
        
    except FileProcessingError:
        raise
    except Exception as e:
        raise FileProcessingError(f"DOCX 파싱 중 오류 발생: {str(e)}")


def extract_text_from_file(file_obj):
    """
    파일 객체에서 텍스트 추출 (통합 함수, OCR 최적화)
    
    Args:
        file_obj: Flask 파일 객체
        
    Returns:
        str: 추출된 텍스트
        
    Raises:
        FileProcessingError: 파일 처리 실패 시
    """
    if not file_obj or not file_obj.filename:
        raise FileProcessingError("유효하지 않은 파일입니다.")
    
    # 파일 타입 검증
    validate_file_type(file_obj.filename)
    
    # 파일 확장자에 따른 파싱 (OCR 비용 최적화)
    filename = file_obj.filename.lower()
    
    if filename.endswith('.pdf'):
        return parse_pdf(file_obj.stream)
    elif filename.endswith('.docx'):
        return parse_docx(file_obj.stream)
    else:
        raise FileProcessingError(f"지원하지 않는 파일 형식입니다: {filename}")


def get_file_info(file_obj):
    """
    파일 정보 추출
    
    Args:
        file_obj: Flask 파일 객체
        
    Returns:
        dict: 파일 정보 (이름, 크기, 타입 등)
    """
    if not file_obj or not file_obj.filename:
        return None
    
    # 파일 크기 계산
    file_obj.stream.seek(0, 2)  # 파일 끝으로 이동
    file_size = file_obj.stream.tell()
    file_obj.stream.seek(0)  # 처음으로 되돌리기
    
    return {
        'filename': file_obj.filename,
        'size': file_size,
        'size_mb': round(file_size / 1024 / 1024, 2),
        'extension': os.path.splitext(file_obj.filename)[1].lower(),
        'is_valid_type': os.path.splitext(file_obj.filename)[1].lower() in ALLOWED_EXTENSIONS,
        'is_valid_size': file_size <= MAX_FILE_SIZE
    } 